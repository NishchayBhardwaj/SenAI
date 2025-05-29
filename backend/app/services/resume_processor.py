import os
import logging
import pytesseract
from PIL import Image
import pdf2image
import spacy
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import re
import io
import tempfile
from groq import Groq
import PyPDF2
import docx
from config.settings import GROQ_API_KEY
from typing import Dict, Any, Optional, List
import json

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Groq client
groq_client = Groq(api_key=GROQ_API_KEY)

class ResumeProcessingError(Exception):
    """Custom exception for resume processing errors"""
    pass

def process_pdf_content(file: io.BytesIO) -> str:
    """Extracts text from a PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        content = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_content = page.extract_text()
            content += page_content
        
        # Check if text extraction failed or returned very little text
        if len(content.strip()) < 100:
            logger.warning("Standard text extraction yielded minimal results. Attempting OCR...")
            content = process_pdf_with_ocr(file)
        
        return content
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        raise ResumeProcessingError(f"Failed to process PDF: {str(e)}")

def process_pdf_with_ocr(file: io.BytesIO) -> str:
    """Extracts text from images/scanned PDFs using OCR."""
    content = ""
    with tempfile.TemporaryDirectory() as path:
        pdf_path = os.path.join(path, "temp.pdf")
        
        # Save the uploaded file
        with open(pdf_path, "wb") as f:
            file.seek(0)
            f.write(file.read())
        file.seek(0)  # Reset file pointer
        
        try:
            images = pdf2image.convert_from_path(pdf_path)
            for image in images:
                content += pytesseract.image_to_string(image)
        except Exception as e:
            logger.error(f"OCR error: {str(e)}")
            raise ResumeProcessingError(f"Failed to perform OCR: {str(e)}")
    
    return content

def process_docx_content(file: io.BytesIO) -> str:
    """Extracts text from a Word document."""
    try:
        doc = docx.Document(file)
        full_content = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_content.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_content.append(cell.text)
        
        extracted_content = '\n'.join(full_content)
        
        # Check if extracted text is minimal
        if len(extracted_content.strip()) < 100:
            logger.warning("Standard text extraction yielded minimal results from DOCX. Attempting OCR on document images...")
            extracted_content = process_docx_images(file) or extracted_content
        
        return extracted_content
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        raise ResumeProcessingError(f"Failed to process DOCX: {str(e)}")

def process_docx_images(file: io.BytesIO) -> str:
    """Extract text from images embedded in a DOCX file using OCR."""
    try:
        # Save uploaded file to temporary directory
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, "temp.docx")
            
            # Save the uploaded file
            with open(temp_file_path, "wb") as f:
                file.seek(0)
                f.write(file.read())
            file.seek(0)  # Reset file pointer
            
            # Load the document again
            doc = docx.Document(temp_file_path)
            
            # Extract and process images
            extracted_content = ""
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Create a PIL Image from binary data
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Use OCR to extract text
                        image_content = pytesseract.image_to_string(image)
                        if image_content.strip():
                            extracted_content += image_content + "\n\n"
                    except Exception as e:
                        logger.error(f"Error processing image in DOCX: {str(e)}")
                        continue
            
            return extracted_content
    except Exception as e:
        logger.error(f"Error extracting images from DOCX: {str(e)}")
        raise ResumeProcessingError(f"Failed to extract images from DOCX: {str(e)}")

def process_txt_content(file: io.BytesIO) -> str:
    """Extracts text from a text file."""
    try:
        content = file.read().decode('utf-8')
        return content
    except Exception as e:
        logger.error(f"Error processing TXT: {str(e)}")
        raise ResumeProcessingError(f"Failed to process TXT: {str(e)}")

def process_file_content(file: io.BytesIO, file_type: str) -> str:
    """Extract text from different file types and validate resume content."""
    try:
        if file_type == "pdf":
            content = process_pdf_content(file)
        elif file_type == "docx":
            content = process_docx_content(file)
        elif file_type == "txt":
            content = process_txt_content(file)
        else:
            raise ResumeProcessingError(f"Unsupported file type: {file_type}")

        # Resume validation logic
        # Consider content invalid if it's blank, too short, or mostly non-alphabetic
        if not content or len(content.strip()) < 50:
            raise ResumeProcessingError("Invalid file, please upload a valid resume (file is empty or too short).")
        # Check if content is mostly non-alphabetic (trash file)
        alpha_chars = sum(c.isalpha() for c in content)
        if alpha_chars < 10 or (alpha_chars / max(1, len(content))) < 0.1:
            raise ResumeProcessingError("Invalid file, please upload a valid resume (file does not contain enough readable text).")

        return content
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        raise ResumeProcessingError(f"Failed to process file: {str(e)}")

def analyze_resume_content(resume_content: str) -> Dict[str, Any]:
    """Extracts structured data from resume text using Groq API."""
    prompt = f"""Extract ONLY the following information from the resume text provided below:
    - Full Name
    - Email Address
    - Phone Number
    - Location (City, State/Country)
    - Education (Degree, Institution, Year) - list all
    - Work Experience (Company, Position, Duration) - list all
    - Skills (Technical and Soft Skills) - list all (only names like python,nextjs,leadership,etc)
    - Years of Experience - IMPORTANT: If not explicitly stated, calculate this by adding up all work experience durations or estimate based on career progression just show the number no explaination needed

    Resume Text:
    {resume_content}

    Return ONLY the extracted information in this exact format - do not include any additional information, analysis, or commentary:

    ## Full Name
    [Extracted name]

    ## Email Address
    [Extracted email]

    ## Phone Number
    [Extracted phone]

    ## Location
    [Extracted location]

    ## Education
    - [Degree], [Institution], [Year]
    - [Additional education entries]

    ## Work Experience
    - [Company], [Position], [Duration]
    - [Additional work experience entries]

    ## Skills
    [List of extracted skills]

    ## Years of Experience
    [Number of years] - YOU MUST PROVIDE THIS! Calculate if not explicitly stated in resume aand display only the number

    If a field is not found, indicate "Not found" for that field only.
    """

    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="llama3-70b-8192",
            temperature=0.2,
            max_tokens=1000
        )
        
        # Parse the response into a structured format
        response_text = chat_completion.choices[0].message.content
        structured_data = {
            'Full Name': '',
            'Email Address': '',
            'Phone Number': '',
            'Location': '',
            'Education': [],
            'Work Experience': [],
            'Skills': [],
            'Years of Experience': 0
        }
        
        try:
            # Split the response into sections
            sections = response_text.split('##')
            
            for section in sections:
                if not section.strip():
                    continue
                    
                # Split section into title and content
                lines = section.strip().split('\n')
                title = lines[0].strip()
                content = '\n'.join(lines[1:]).strip()
                
                if 'Full Name' in title:
                    structured_data['Full Name'] = content
                elif 'Email Address' in title:
                    structured_data['Email Address'] = content
                elif 'Phone Number' in title:
                    structured_data['Phone Number'] = content
                elif 'Location' in title:
                    structured_data['Location'] = content
                elif 'Education' in title:
                    # Parse education entries
                    entries = [entry.strip('- ').strip() for entry in content.split('\n') if entry.strip()]
                    for entry in entries:
                        if entry and entry != 'Not found':
                            parts = [part.strip() for part in entry.split(',')]
                            if len(parts) >= 3:
                                # Extract year from the last part, handling cases where it might be mixed with location
                                year_part = parts[-1].strip()
                                # Try to extract a 4-digit year
                                year_match = re.search(r'\b(19|20)\d{2}\b', year_part)
                                year = year_match.group(0) if year_match else None
                                
                                structured_data['Education'].append({
                                    'degree': parts[0].strip(),
                                    'institution': parts[1].strip(),
                                    'year': year
                                })
                elif 'Work Experience' in title:
                    # Parse work experience entries
                    entries = [entry.strip('- ').strip() for entry in content.split('\n') if entry.strip()]
                    structured_data['Work Experience'] = entries
                elif 'Skills' in title:
                    # Parse skills
                    skills = [skill.strip() for skill in content.split(',') if skill.strip()]
                    structured_data['Skills'] = skills
                elif 'Years of Experience' in title:
                    # Parse years of experience
                    try:
                        years = int(content.strip())
                        structured_data['Years of Experience'] = years
                    except ValueError:
                        structured_data['Years of Experience'] = 0
            
            return structured_data
        except Exception as e:
            logger.error(f"Error parsing resume content: {str(e)}")
            # Return basic structure with the raw text as full name if parsing fails
            return {
                'Full Name': response_text,
                'Email Address': '',
                'Phone Number': '',
                'Location': '',
                'Education': [],
                'Work Experience': [],
                'Skills': [],
                'Years of Experience': 0
            }
    except Exception as e:
        logger.error(f"Error analyzing resume content: {str(e)}")
        raise ResumeProcessingError(f"Failed to analyze resume content: {str(e)}") 
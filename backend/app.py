import os
import streamlit as st
import pytesseract
from PIL import Image
import pdf2image
import spacy
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import re
from dotenv import load_dotenv
import groq
import json
import sys
import io
import tempfile
from groq import Groq
import PyPDF2
import docx
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configure pytesseract path for Windows
if sys.platform.startswith('win'):
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Initialize NLTK
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
except Exception as e:
    st.error(f"Error downloading NLTK data: {str(e)}")

# Initialize spaCy
try:
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    st.error(f"Error loading spaCy model: {str(e)}")
    st.info("Please run: python -m spacy download en_core_web_sm")

# Initialize Groq client
try:
    groq_client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
except Exception as e:
    st.error(f"Error initializing Groq client: {str(e)}")
    st.info("Please check your GROQ_API_KEY in .env file")

def check_poppler():
    """Check if Poppler is installed and accessible"""
    try:
        # Try to convert a simple PDF to check Poppler installation
        pdf2image.convert_from_path("test.pdf", first_page=1, last_page=1)
        return True
    except Exception as e:
        if "poppler" in str(e).lower():
            st.error("""
            Poppler is not installed or not in PATH. Please follow these steps:
            1. Download Poppler from: https://github.com/oschwartz10612/poppler-windows/releases/
            2. Extract to a permanent location (e.g., C:\\Program Files\\poppler)
            3. Add the bin directory to system PATH (e.g., C:\\Program Files\\poppler\\Library\\bin)
            4. Restart your terminal/IDE
            """)
        return False

def extract_text_from_pdf(pdf_path):
    """Convert PDF to images and extract text using OCR"""
    if not check_poppler():
        return None
        
    try:
        # Convert PDF to images
        images = pdf2image.convert_from_path(pdf_path)
        text = ""
        
        # Extract text from each page
        for image in images:
            text += pytesseract.image_to_string(image)
        
        return text
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
        return None

def extract_text_from_image(image_path):
    """Extract text from image using OCR"""
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error processing image: {str(e)}")
        return None

def process_with_groq(text):
    """Process extracted text with Groq LLM to get structured information"""
    if not os.environ.get("GROQ_API_KEY"):
        st.error("GROQ_API_KEY not found in environment variables")
        return None

    prompt = f"""
    Extract the following information from this resume text:
    - Full Name
    - Email Address
    - Phone Number
    - Location (City, State/Country)
    - Education (Degree, Institution, Year)
    - Work Experience (Company, Position, Duration)
    - Skills (Technical and Soft Skills)
    - Years of Experience

    Resume text:
    {text}

    Return the information in JSON format.
    """

    try:
        response = groq_client.chat.completions.create(
            model="llama2-70b-4096",
            messages=[
                {"role": "system", "content": "You are a resume parser that extracts structured information."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=1000
        )
        
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        st.error(f"Error processing with Groq: {str(e)}")
        return None

def process_pdf_content(file):
    """Extracts text from a PDF file."""
    pdf_reader = PyPDF2.PdfReader(file)
    content = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        page_content = page.extract_text()
        content += page_content
    
    # Check if text extraction failed or returned very little text
    if len(content.strip()) < 100:  # Adjust threshold as needed
        st.warning("Standard text extraction yielded minimal results. Attempting OCR...")
        content = process_pdf_with_ocr(file)
    
    return content

def process_pdf_with_ocr(file):
    """Extracts text from images/scanned PDFs using OCR."""
    content = ""
    with tempfile.TemporaryDirectory() as path:
        pdf_path = os.path.join(path, "temp.pdf")
        
        # Save the uploaded file
        with open(pdf_path, "wb") as f:
            file.seek(0)
            f.write(file.read())
        file.seek(0)  # Reset file pointer for future operations
        
        try:
            images = pdf2image.convert_from_path(pdf_path)
            for image in images:
                content += pytesseract.image_to_string(image)
        except Exception as e:
            st.error(f"OCR error: {str(e)}")
    
    return content

def process_docx_content(file):
    """Extracts text from a Word document."""
    doc = docx.Document(file)
    full_content = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_content.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_content.append(cell.text)
    
    extracted_content = '\n'.join(full_content)
    
    # Check if extracted text is minimal
    if len(extracted_content.strip()) < 100:  # Adjust threshold as needed
        st.warning("Standard text extraction yielded minimal results from DOCX. Attempting OCR on document images...")
        extracted_content = process_docx_images(file) or extracted_content
    
    return extracted_content

def process_docx_images(file):
    """Extract text from images embedded in a DOCX file using OCR."""
    try:
        # Save uploaded file to temporary directory
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, "temp.docx")
            
            # Save the uploaded file
            with open(temp_file_path, "wb") as f:
                file.seek(0)
                f.write(file.read())
            file.seek(0)  # Reset file pointer
            
            # Load the document again
            doc = docx.Document(temp_file_path)
            
            # Extract and process images
            extracted_content = ""
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Create a PIL Image from binary data
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Use OCR to extract text
                        image_content = pytesseract.image_to_string(image)
                        if image_content.strip():
                            extracted_content += image_content + "\n\n"
                    except Exception as e:
                        st.error(f"Error processing image in DOCX: {str(e)}")
                        continue
            
            return extracted_content
    except Exception as e:
        st.error(f"Error extracting images from DOCX: {str(e)}")
        return ""

def process_txt_content(file):
    """Extracts text from a text file."""
    content = file.read().decode('utf-8')
    return content

def process_file_content(file, file_type):
    """Extract text from different file types."""
    if file_type == "pdf":
        return process_pdf_content(file)
    elif file_type == "docx":
        return process_docx_content(file)
    elif file_type == "txt":
        return process_txt_content(file)
    else:
        return "Unsupported file type"

def analyze_resume_content(resume_content):
    """Extracts structured data from resume text using Groq API."""
    prompt = f"""Extract ONLY the following information from the resume text provided below:
    - Full Name
    - Email Address
    - Phone Number
    - Location (City, State/Country)
    - Education (Degree, Institution, Year) - list all
    - Work Experience (Company, Position, Duration) - list all
    - Skills (Technical and Soft Skills) - list all (only names like python,nextjs,leadership,etc)
    - Years of Experience - IMPORTANT: If not explicitly stated, calculate this by adding up all work experience durations or estimate based on career progression just show the number no explaination needed

    Resume Text:
    {resume_content}

    Return ONLY the extracted information in this exact format - do not include any additional information, analysis, or commentary:

    ## Full Name
    [Extracted name]

    ## Email Address
    [Extracted email]

    ## Phone Number
    [Extracted phone]

    ## Location
    [Extracted location]

    ## Education
    - [Degree], [Institution], [Year]
    - [Additional education entries]

    ## Work Experience
    - [Company], [Position], [Duration]
    - [Additional work experience entries]

    ## Skills
    [List of extracted skills]

    ## Years of Experience
    [Number of years] - YOU MUST PROVIDE THIS! Calculate if not explicitly stated in resume aand display only the number

    If a field is not found, indicate "Not found" for that field only.
    """

    chat_completion = groq_client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="llama3-70b-8192", # Using LLaMA 3 70B model
        temperature=0.2, # Lower temperature for more consistent and precise output
        max_tokens=1000 # Limit response length to avoid unnecessary content
    )
    
    # Parse the response into a structured format
    response_text = chat_completion.choices[0].message.content
    structured_data = {
        'Full Name': '',
        'Email Address': '',
        'Phone Number': '',
        'Location': '',
        'Education': [],
        'Work Experience': [],
        'Skills': [],
        'Years of Experience': 0
    }
    
    try:
        # Split the response into sections
        sections = response_text.split('##')
        
        for section in sections:
            if not section.strip():
                continue
                
            # Split section into title and content
            lines = section.strip().split('\n')
            title = lines[0].strip()
            content = '\n'.join(lines[1:]).strip()
            
            if 'Full Name' in title:
                structured_data['Full Name'] = content
            elif 'Email Address' in title:
                structured_data['Email Address'] = content
            elif 'Phone Number' in title:
                structured_data['Phone Number'] = content
            elif 'Location' in title:
                structured_data['Location'] = content
            elif 'Education' in title:
                # Parse education entries
                entries = [entry.strip('- ').strip() for entry in content.split('\n') if entry.strip()]
                for entry in entries:
                    if entry and entry != 'Not found':
                        parts = [part.strip() for part in entry.split(',')]
                        if len(parts) >= 3:
                            # Extract year from the last part, handling cases where it might be mixed with location
                            year_part = parts[-1].strip()
                            # Try to extract a 4-digit year
                            year_match = re.search(r'\b(19|20)\d{2}\b', year_part)
                            year = year_match.group(0) if year_match else None
                            
                            structured_data['Education'].append({
                                'degree': parts[0].strip(),
                                'institution': parts[1].strip(),
                                'year': year
                            })
            elif 'Work Experience' in title:
                # Parse work experience entries
                entries = [entry.strip('- ').strip() for entry in content.split('\n') if entry.strip()]
                structured_data['Work Experience'] = entries
            elif 'Skills' in title:
                # Parse skills
                skills = [skill.strip() for skill in content.split(',') if skill.strip()]
                structured_data['Skills'] = skills
            elif 'Years of Experience' in title:
                # Parse years of experience
                try:
                    years = int(content.strip())
                    structured_data['Years of Experience'] = years
                except ValueError:
                    structured_data['Years of Experience'] = 0
        
        return structured_data
    except Exception as e:
        logger.error(f"Error parsing resume content: {str(e)}")
        # Return basic structure with the raw text as full name if parsing fails
        return {
            'Full Name': response_text,
            'Email Address': '',
            'Phone Number': '',
            'Location': '',
            'Education': [],
            'Work Experience': [],
            'Skills': [],
            'Years of Experience': 0
        }

def main():
    st.title("Resume Processing System")
    st.write("Upload a resume (PDF, DOCX, or TXT) to extract information")

    # Check for required dependencies
    if not os.path.exists(pytesseract.pytesseract.tesseract_cmd):
        st.error("""
        Tesseract OCR is not installed or not found. Please:
        1. Download Tesseract from: https://github.com/UB-Mannheim/tesseract/wiki
        2. Install it to the default location (C:\\Program Files\\Tesseract-OCR)
        3. Restart your terminal/IDE
        """)
        return

    uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx', 'txt'])

    if uploaded_file is not None:
        st.write("File Uploaded Successfully!")
        
        # Get file type
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        # Extract text based on file type
        with st.spinner(f"Extracting text from {file_type.upper()} file..."):
            resume_content = process_file_content(uploaded_file, file_type)
        st.subheader("Extracted Text:")
        st.text_area("Resume Content", resume_content, height=300)

        # Process with Groq
        if st.button("Extract Information"):
            if not os.environ.get("GROQ_API_KEY") or os.environ.get("GROQ_API_KEY") == "YOUR_GROQ_API_KEY":
                st.error("GROQ_API_KEY not configured. Please set it in the .env file.")
            elif not resume_content.strip():
                st.warning("The extracted text is empty. Cannot process.")
            else:
                with st.spinner("Extracting information using Groq LLaMA-3..."):
                    extracted_data = analyze_resume_content(resume_content)
                st.subheader("Extracted Information:")
                st.markdown(extracted_data)

if __name__ == "__main__":
    main()